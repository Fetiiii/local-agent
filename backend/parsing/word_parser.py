from __future__ import annotations

from pathlib import Path
from typing import Dict, Any

from docx import Document  # type: ignore

from backend.utils.validators import ensure_file_exists, validate_extension
from backend.parsing.text_cleaner import clean_text, excerpt


def parse_word(
    path: str | Path,
    max_paragraphs: int = 200,
    max_chars: int = 8000,
) -> Dict[str, Any]:
    """
    Extract paragraph and simple table text from a Word document.
    Limits paragraphs/characters to stay lightweight.
    """
    if max_paragraphs <= 0:
        raise ValueError("max_paragraphs must be positive")
    if max_chars <= 0:
        raise ValueError("max_chars must be positive")

    validate_extension(path, {".docx"})
    doc_path = ensure_file_exists(path)
    doc = Document(str(doc_path))

    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))

    combined = paragraphs + table_rows
    combined = combined[:max_paragraphs]

    content = clean_text("\n".join(combined))
    if len(content) > max_chars:
        content = content[:max_chars].rstrip()

    return {
        "path": str(doc_path),
        "content": content,
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "preview": excerpt(content, max_chars=1200),
    }
