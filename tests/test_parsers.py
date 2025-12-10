import pandas as pd
from docx import Document

from backend.parsing import (
    normalize_whitespace,
    clean_text,
    excerpt,
    parse_excel,
    parse_word,
)
import backend.parsing.pdf_parser as pdf_parser


def test_text_cleaner_collapses_spaces_and_tags():
    raw = "a   b \n\n<|analysis|> c"
    assert normalize_whitespace(raw) == "a b\n\nc"


def test_clean_text_no_newlines():
    assert clean_text(" a \r\n b\t\t", keep_newlines=False) == "a b"


def test_excerpt_shortens():
    assert excerpt("abcd", max_chars=4) == "a..."


def test_parse_excel_preview(tmp_path):
    df = pd.DataFrame({"col1": [1, 2], "col2": ["x", "y"]})
    path = tmp_path / "sample.xlsx"
    df.to_excel(path, index=False)

    res = parse_excel(path, max_rows=1, max_cols=1)
    assert res["sheet"] == "Sheet1"
    assert res["row_count_preview"] == 1
    assert res["col_count_preview"] == 1
    assert "col1" in res["columns"][0]
    assert "1" in res["preview_csv"]


def test_parse_word_reads_tables(tmp_path):
    doc = Document()
    doc.add_paragraph("hello world")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"

    path = tmp_path / "sample.docx"
    doc.save(path)

    res = parse_word(path, max_paragraphs=5, max_chars=100)
    assert "hello world" in res["content"]
    assert "A | B" in res["content"]
    assert res["paragraphs"] == 1
    assert res["tables"] == 1


def test_parse_pdf_respects_max_pages(monkeypatch, tmp_path):
    path = tmp_path / "sample.pdf"
    path.write_bytes(b"%PDF-1.4")

    class DummyPage:
        def __init__(self, text):
            self.text = text

        def extract_text(self):
            return self.text

    class DummyReader:
        def __init__(self, *_args, **_kwargs):
            self.pages = [DummyPage("first page"), DummyPage("second page")]

    monkeypatch.setattr(pdf_parser, "PdfReader", DummyReader)

    res = pdf_parser.parse_pdf(path, max_pages=1)
    assert res["pages_read"] == 1
    assert "first page" in res["content"]
    assert "second page" not in res["content"]
